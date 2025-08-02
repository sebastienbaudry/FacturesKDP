#!/usr/bin/env python3
"""
Générateur de factures Word et PDF automatisé pour les revenus Amazon KDP
Auteur: Sébastien Baudry - assisté de Claude 4 Sonnet, Gemini Pro 2.5, ChatGPT 4o
Version: 3.2 – correction détails + ajout format PDF
"""

import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import calendar
import locale
import json
from pathlib import Path
from fpdf import FPDF
from fpdf.enums import XPos, YPos
import sys

# ---------- CONFIG / UTILITAIRES ---------------------------------------------
def charger_configuration(chemin_config="config.json"):
    try:
        config_path = Path(chemin_config)
        if not config_path.is_file():
            return None, f"ERREUR: Fichier de configuration '{chemin_config}' non trouvé."
        with config_path.open('r', encoding='utf-8') as f:
            config = json.load(f)
        champs_obligatoires = [
            ('entreprise', 'nom'), ('entreprise', 'adresse'), ('entreprise', 'siret'),
            ('entreprise', 'tva_intra'), ('entreprise', 'iban'), ('entreprise', 'bic')
        ]
        for section, champ in champs_obligatoires:
            valeur = config.get(section, {}).get(champ, "")
            if not valeur or '[' in str(valeur):
                return None, f"ERREUR: Le champ '{section}.{champ}' n'est pas configuré dans {chemin_config}"
        return config, "Configuration chargée."
    except Exception as e:
        return None, f"ERREUR lors du chargement de la configuration: {e}"

def setup_locale():
    # Tenter de définir la locale sur le français.
    # Les noms de locale peuvent varier légèrement selon le système d'exploitation.
    # 'fr_FR.UTF-8' est courant sur Linux/macOS.
    # 'fra_FRA.1252' ou 'French_France.1252' est courant sur Windows.
    # On essaie plusieurs options.
    try:
        locale.setlocale(locale.LC_TIME, 'fr_FR.UTF-8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'fra_FRA.1252') # Pour Windows
        except locale.Error:
            locale.setlocale(locale.LC_TIME, 'French_France.1252') # Autre pour Windows

# ---------- LECTURE DES DONNÉES ----------------------------------------------
def lire_fichier_kdp(chemin_fichier):
    try:
        excel_path = Path(chemin_fichier)
        if not excel_path.is_file():
            return None, f"ERREUR: Le fichier Excel '{chemin_fichier}' est introuvable."
        df = pd.read_excel(excel_path, sheet_name='Paiements')
        df.columns = df.columns.str.strip()
        colonnes_requises = [
            'Période de vente - Date de début', 'Marché', 'Numéro de paiement',
            'Devise', 'Redevance accumulée', 'Montant du paiement'
        ]
        for col in colonnes_requises:
            if col not in df.columns:
                return None, f"ERREUR: La colonne '{col}' est manquante."
        return df, f"Fichier lu avec succès: {len(df)} lignes."
    except Exception as e:
        return None, f"Erreur lors de la lecture du fichier Excel: {e}"

def extraire_donnees_periode(df, annee, mois):
    """
    Extrait les données pour une période donnée (année/mois).
    Inclut les lignes principales et leurs détails même sans date.
    """
    periode_debut = f"{annee}-{mois:02d}-01"
    mois_str = f"{annee}-{mois:02d}"

    # Filtre : lignes principales OU lignes dont la colonne 'Détail' contient le mois ciblé
    mask_periode = (df['Période de vente - Date de début'] == periode_debut) | \
                   (df['Détail'].astype(str).str.startswith(mois_str))

    donnees_periode = df[mask_periode].copy()

    if donnees_periode.empty:
        return None, f"Aucune donnée trouvée pour {calendar.month_name[mois]} {annee}"

    return donnees_periode, f"Données trouvées : {len(donnees_periode)} lignes."

def regrouper_par_marche(donnees):
    """
    Regroupe les données par marché (lignes principales + détails)
    """
    marches_data = {}

    # 1) Lignes principales (avec numéro de paiement)
    principales = donnees[donnees['Numéro de paiement'].notna()].copy()
    principales['Redevance accumulée'] = pd.to_numeric(
        principales['Redevance accumulée'], errors='coerce').fillna(0)
    principales['Montant du paiement'] = pd.to_numeric(
        principales['Montant du paiement'], errors='coerce').fillna(0)

    marche_par_idx = {}
    for idx, ligne in principales.iterrows():
        marche = ligne['Marché']
        if marche not in marches_data:
            marches_data[marche] = {
                'devise_origine': ligne['Devise'],
                'taux_change': ligne.get('Taux de change'),
                'total_origine': 0.0,
                'total_eur': 0.0,
                'details': []
            }
        marches_data[marche]['total_origine'] += ligne['Redevance accumulée']
        marches_data[marche]['total_eur'] += ligne['Montant du paiement']
        marche_par_idx[idx] = marche

    # 2) Détails (lignes sans numéro de paiement)
    marche_actuel = None
    for idx, ligne in donnees.iterrows():
        if idx in marche_par_idx:
            marche_actuel = marche_par_idx[idx]
            continue
        if marche_actuel and pd.notna(ligne.get('Redevance accumulée')):
            col = 'Source' if 'Source' in donnees.columns else 'Détail'
            designation = ligne.get(col, '') or 'Redevance KDP'
            devise = ligne.get('Devise') or marches_data[marche_actuel]['devise_origine']
            montant = float(ligne['Redevance accumulée'])
            marches_data[marche_actuel]['details'].append({
                'designation': designation,
                'devise': devise,
                'montant': montant
            })

    return marches_data, f"Marchés trouvés: {list(marches_data.keys())}"




# ---------- GÉNÉRATION WORD --------------------------------------------------
def generer_numero_facture(config, annee, mois, numero_personnalise=None):
    if numero_personnalise:
        return numero_personnalise
    fmt = config['facture'].get('format_numero', "{annee}-{mois:02d}-01")
    prefixe = config['facture'].get('prefixe_numero', "FACT")
    num = fmt.format(annee=annee, mois=mois)
    return f"{prefixe}-{num}" if prefixe else num

def obtenir_date_paiement(config, date_personnalisee=None):
    return date_personnalisee or config['facture'].get('date_paiement_defaut', "Non spécifiée")

def _cell_bold(cell):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if not p.runs:
        p.add_run('')
    p.runs[0].bold = True

def creer_facture_word(marches_data, annee, mois, config, numero_facture=None, date_paiement=None):
    setup_locale()
    doc = Document()
    for s in doc.sections:
        s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = (Inches(i) for i in (.5,.5,.8,.8))

    ent = config['entreprise']
    cli = config['client']
    num = generer_numero_facture(config, annee, mois, numero_facture)
    date_pmt = obtenir_date_paiement(config, date_paiement)
    nom_mois = calendar.month_name[mois]

    # En-tête
    p = doc.add_paragraph()
    p.add_run(ent['nom']).bold = True
    p.add_run(f"\n{ent['adresse']}\nSIRET : {ent['siret']}\nTVA intracommunautaire : {ent['tva_intra']}")
    if ent.get('code_ape'):
        p.add_run(f"\nCode APE : {ent['code_ape']}")
    if ent.get('forme_juridique'):
        p.add_run(f"\nForme juridique : {ent['forme_juridique']}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Destinataire de la facture").bold = True
    p.add_run(f"\n\n{cli['nom']}\n{cli['adresse']}\nTVA intracommunautaire : {cli['tva_intra']}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Facture n° : {num}").bold = True
    p.add_run(f"\nDate de la facture : {datetime.now():%d/%m/%Y}")
    p.add_run(f"\nPériode concernée : Revenus de {nom_mois} {annee}")
    p.add_run(f"\nDate de paiement constatée : {date_pmt}")
    p.add_run(f"\nMode de règlement : {config['facture']['mode_reglement']}")
    p.add_run(f"\nIBAN : {ent['iban']}\nBIC : {ent['bic']}")

    total_eur = 0

    # Pour chaque marché, un tableau distinct
    for marche, data in marches_data.items():
        doc.add_paragraph()
        titre = doc.add_paragraph(f"Détail pour le marché : {marche}")
        titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titre.runs[0].bold = True

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        headers = ['Marché','Désignation','Devise','Montant net','Taux','Montant EUR']
        for i,h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            _cell_bold(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for det in data.get('details', []):
            montant = det['montant']
            devise = det['devise']
            taux = float(data['taux_change']) if data['taux_change'] and data['devise_origine'] != 'EUR' else 1
            montant_eur = montant * taux

            row = table.add_row().cells
            row[0].text = marche
            row[1].text = det['designation']
            row[2].text = devise
            row[3].text = f"{montant:.2f}"
            row[4].text = f"{taux:.3f}" if taux != 1 else ""
            row[5].text = f"{montant_eur:.2f} €"

        # Ligne de total du marché en gras
        row = table.add_row().cells
        for c in row: _cell_bold(c)
        row[0].text = marche
        row[1].text = "TOTAL"
        row[2].text = data['devise_origine']
        row[3].text = f"{data['total_origine']:.2f}"
        if data.get('taux_change') and pd.notna(data['taux_change']):
            row[4].text = str(data['taux_change'])
        row[5].text = f"{data['total_eur']:.2f} €"

        total_eur += data['total_eur']

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Montant total HT : {total_eur:.2f} €\n")
    p.add_run("TVA : 0,00 € (Autoliquidation)\n\n")
    p.add_run(f"Montant TTC : {total_eur:.2f} €").bold = True
    doc.add_paragraph()
    p = doc.add_paragraph(config['messages']['autoliquidation'])
    p.italic = True
    return doc, total_eur


# ---------- GÉNÉRATION PDF ---------------------------------------------------
def creer_facture_pdf(marches_data, annee, mois, config, numero_facture=None, date_paiement=None):
    pdf = FPDF('P','mm','A4')
    pdf.set_auto_page_break(True,15)
    pdf.set_margins(20,20,20)
    pdf.add_page()
    font = 'Helvetica'
    euro = chr(128)

    ent = config['entreprise']
    cli = config['client']
    num = generer_numero_facture(config, annee, mois, numero_facture)
    date_pmt = obtenir_date_paiement(config, date_paiement)
    nom_mois = calendar.month_name[mois]

    # En-tête entreprise
    pdf.set_font(font,'B',12)
    pdf.cell(0,5,ent['nom'],new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.set_font(font,'',10)
    for ln in ent['adresse'].splitlines():
        pdf.cell(0,5,ln.strip(),new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.cell(0,5,f"SIRET : {ent['siret']}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.cell(0,5,f"TVA intracommunautaire : {ent['tva_intra']}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.ln(10)

    # Destinataire
    pdf.set_font(font,'B',11)
    pdf.cell(0,5,"Destinataire de la facture",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.ln(2)
    pdf.set_font(font,'',10)
    pdf.cell(0,5,cli['nom'],new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    for ln in cli['adresse'].splitlines():
        pdf.cell(0,5,ln.strip(),new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.ln(10)

    # Infos facture
    pdf.set_font(font,'B',12)
    pdf.cell(0,5,f"Facture n° : {num}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.set_font(font,'',10)
    pdf.cell(0,5,f"Date : {datetime.now():%d/%m/%Y}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.cell(0,5,f"Période : Revenus de {nom_mois} {annee}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.cell(0,5,f"Date paiement constatée : {date_pmt}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.cell(0,5,f"Mode règlement : {config['facture']['mode_reglement']}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.cell(0,5,f"IBAN : {ent['iban']}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.cell(0,5,f"BIC : {ent['bic']}",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.ln(10)

    total_eur = 0
    w = [25,60,25,20,20,25]  # Largeurs des colonnes

    # Un tableau par marché
    for marche, data in marches_data.items():
        pdf.set_font(font,'B',11)
        pdf.cell(0,7,f"Détail pour le marché : {marche}",0,1,'C')
        pdf.ln(1)

        # En-têtes du tableau
        pdf.set_font(font,'B',8)
        headers = ['Marché','Désignation','Devise','Net','Tx',f'Montant {euro}']
        for i,h in enumerate(headers):
            pdf.cell(w[i],7,h,1,0,'C')
        pdf.ln()

        pdf.set_font(font,'',8)
        for det in data.get('details', []):
            montant = det['montant']
            devise = det['devise']
            taux = float(data['taux_change']) if data['taux_change'] and data['devise_origine'] != 'EUR' else 1
            montant_eur = montant * taux

            pdf.cell(w[0],7,marche,1)
            pdf.cell(w[1],7,det['designation'],1)
            pdf.cell(w[2],7,devise,1)
            pdf.cell(w[3],7,f"{montant:.2f}",1,align='R')
            pdf.cell(w[4],7,f"{taux:.3f}" if taux != 1 else "",1,align='R')
            pdf.cell(w[5],7,f"{montant_eur:.2f} {euro}",1,align='R')
            pdf.ln()

        # Ligne total du marché en gras
        pdf.set_font(font,'B',8)
        pdf.cell(w[0],7,marche,1)
        pdf.cell(w[1],7,"TOTAL",1)
        pdf.cell(w[2],7,data['devise_origine'],1)
        pdf.cell(w[3],7,f"{data['total_origine']:.2f}",1,align='R')
        taux_str = str(data['taux_change']) if pd.notna(data.get('taux_change')) else ""
        pdf.cell(w[4],7,taux_str,1,align='R')
        pdf.cell(w[5],7,f"{data['total_eur']:.2f} {euro}",1,align='R')
        pdf.ln(10)
        pdf.set_font(font,'',8)

        total_eur += data['total_eur']

    # Totaux finaux
    pdf.ln(5)
    pdf.set_font(font,'',10)
    pdf.cell(0,5,f"Total HT : {total_eur:.2f} {euro}",0,1,'R')
    pdf.cell(0,5,f"TVA : 0,00 {euro} (Autoliquidation)",0,1,'R')
    pdf.ln(2)
    pdf.set_font(font,'B',12)
    pdf.cell(0,7,f"Total TTC : {total_eur:.2f} {euro}",0,1,'R')
    pdf.ln(10)
    pdf.set_font(font,'I',9)
    pdf.multi_cell(0,5,config['messages']['autoliquidation'])

    return pdf, total_eur


# ---------- GÉNÉRATION / SAVE -------------------------------------------------
def generer_nom_fichier_sortie(config, annee, mois, nom_personnalise=None, extension=".docx"):
    if nom_personnalise:
        return str(Path(nom_personnalise).with_suffix(extension))
    fmt = config['fichiers']['format_nom_sortie']
    dossier = Path(config['fichiers'].get('dossier_sortie','.'))
    dossier.mkdir(parents=True, exist_ok=True)
    return dossier / Path(fmt.format(annee=annee, mois=mois)).with_suffix(extension)

def generer_facture_logic(fichier_excel, annee, mois, format_sortie,
                          config_path='config.json', numero_facture=None, date_paiement=None):
    logs = []
    fichiers = []

    # Charger la configuration
    config, msg = charger_configuration(config_path)
    logs.append(msg)
    if not config:
        return False, "\n".join(logs), []

    # Charger le fichier Excel
    df, msg = lire_fichier_kdp(fichier_excel)
    logs.append(msg)
    if df is None:
        return False, "\n".join(logs), []

    # Extraire les données de la période
    donnees, msg = extraire_donnees_periode(df, annee, mois)
    logs.append(msg)
    if donnees is None:
        return False, "\n".join(logs), []

    # Regrouper par marché
    marches_data, msg = regrouper_par_marche(donnees)
    logs.append(msg)
    if not marches_data:
        logs.append("Aucune donnée de revenus regroupée.")
        return False, "\n".join(logs), []

    total = 0
    try:
        if format_sortie in ('docx', 'both'):
            doc, total = creer_facture_word(marches_data, annee, mois, config, numero_facture, date_paiement)
            nom_docx = generer_nom_fichier_sortie(config, annee, mois, extension=".docx")
            doc.save(nom_docx)
            fichiers.append(str(nom_docx))
            logs.append(f"✅ DOCX : {nom_docx}")

        if format_sortie in ('pdf', 'both'):
            pdf, total = creer_facture_pdf(marches_data, annee, mois, config, numero_facture, date_paiement)
            nom_pdf = generer_nom_fichier_sortie(config, annee, mois, extension=".pdf")
            pdf.output(nom_pdf)
            fichiers.append(str(nom_pdf))
            logs.append(f"✅ PDF : {nom_pdf}")

    except Exception as e:
        logs.append(f"❌ Erreur : {e}")
        return False, "\n".join(logs), []

    logs.append("-" * 50)
    logs.append(f"Montant total : {total:.2f} €")
    logs.append("🎉 Terminé ! Fichiers : " + ", ".join(fichiers))

    return True, "\n".join(logs), fichiers


# ---------- MAIN CLI (optionnel) ---------------------------------------------
if __name__ == "__main__":
    print("Utilisez generateur_factures_kdp.py pour l’interface graphique.")
